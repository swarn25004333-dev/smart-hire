"""Resume parsing service.

Extracts raw text from PDF (PyMuPDF) and DOCX (python-docx) files and
performs a heuristic structure pass. The optional AI layer can later
refine the structured output.
"""

import io
import re
from datetime import datetime
from typing import Dict, List, Optional

from app.models.schemas import Education, Experience, ParsedResume, Project

# ---------------------------------------------------------------------------
# Skill keyword lexicon (used by the offline heuristic engine)
# ---------------------------------------------------------------------------
SKILL_LEXICON: List[str] = [
    # Languages
    "python", "javascript", "typescript", "java", "c++", "c#", "ruby", "go",
    "golang", "rust", "kotlin", "swift", "scala", "php", "sql", "r", "matlab",
    "bash", "powershell", "html", "css",
    # ML / AI / Data
    "machine learning", "deep learning", "tensorflow", "pytorch", "keras",
    "scikit-learn", "pandas", "numpy", "nlp", "computer vision", "openai",
    "langchain", "llm", "generative ai", "reinforcement learning", "spark",
    "hadoop", "kafka", "airflow", "mlflow", "data science", "data analysis",
    "statistics", "bigquery", "tableau", "power bi", "looker",
    # Cloud / DevOps
    "aws", "azure", "gcp", "google cloud", "docker", "kubernetes", "k8s",
    "terraform", "jenkins", "ci/cd", "git", "github", "gitlab", "linux",
    "nginx", "serverless", "lambda", "s3", "ec2", "sagemaker", "mlops",
    # Web / Backend / Frontend
    "react", "react native", "node.js", "nodejs", "express", "django",
    "flask", "fastapi", "spring", "spring boot", "next.js", "vue", "angular",
    "tailwind", "rest api", "graphql", "websockets", "microservices",
    "redis", "elasticsearch", "rabbitmq",
    # Databases
    "postgresql", "mysql", "mongodb", "sqlite", "oracle", "cassandra",
    "dynamodb", "firebase", "supabase",
    # Mobile / Other
    "flutter", "dart", "android", "ios", "swiftui", "unity", "unreal",
    "agile", "scrum", "kanban", "jira", "confluence", "sap",
]

EDUCATION_KEYWORDS: List[str] = [
    "bachelor", "master", "ph.d", "phd", "mba", "b.tech", "b.e.", "m.tech",
    "m.sc", "b.sc", "b.a.", "m.c.a", "b.c.a", "associate", "diploma",
    "high school", "university", "college", "institute", "degree",
]

CERTIFICATION_KEYWORDS: List[str] = [
    "aws certified", "azure certified", "gcp certified", "google certified",
    "pmp", "cissp", "scrum master", "comptia", "tensorflow developer",
    "machine learning specialization", "coursera", "udemy", "databricks",
    "kubernetes (ckad)", "cka", "oracle certified", "microsoft certified",
    "red hat certified", "certified kubernetes", "aws solutions architect",
    "aws developer", "aws machine learning", "a+", "network+", "security+",
]

# Maps a section heading to its canonical category.
SECTION_ALIASES = {
    "summary": "summary", "profile": "summary", "objective": "summary",
    "about me": "summary", "career objective": "summary",
    "education": "education", "academics": "education",
    "academic background": "education", "academic": "education",
    "skills": "skills", "technical skills": "skills",
    "core competencies": "skills", "competencies": "skills",
    "technologies": "skills", "tech skills": "skills",
    "experience": "experience", "work experience": "experience",
    "professional experience": "experience", "work history": "experience",
    "employment history": "experience", "employment": "experience",
    "professional background": "experience",
    "projects": "projects", "project experience": "projects",
    "personal projects": "projects", "academic projects": "projects",
    "key projects": "projects",
    "certifications": "certifications", "certificates": "certifications",
    "certification": "certifications", "licenses": "certifications",
    "licences": "certifications", "licenses & certifications": "certifications",
    "achievements": "achievements", "awards": "achievements",
    "awards & achievements": "achievements", "honors": "achievements",
    "honours": "achievements", "publications": "achievements",
}

NAME_STOPWORDS = {
    "resume", "curriculum", "vitae", "cv", "profile", "summary", "objective",
    "contact", "education", "skills", "experience", "projects", "name",
}

EMAIL_RE = re.compile(r"[\w.+-]+@[\w-]+\.[\w.]+")
PHONE_RE = re.compile(r"(?:\+?\d{1,3}[-.\s]?)?(?:\(?\d{2,4}\)?[-.\s]?)?\d{3}[-.\s]?\d{3,4}[-.\s]?\d{2,4}")
YEARS_RE = re.compile(r"(\d+(?:\.\d+)?)\s*(?:\+|to|-)?\s*(?:years?|yrs|yr)")
DATE_RANGE_RE = re.compile(
    r"^\s*(?:(?:19|20)\d{2})\s*[-–—to]\s*(?:(?:19|20)\d{2}|present|current|now)\s*$",
    re.IGNORECASE,
)

_SKILL_PATTERN_CACHE: Dict[str, "re.Pattern"] = {}


def _matches(text: str, skill: str) -> bool:
    """Word-boundary aware substring match for a skill token."""
    pattern = _SKILL_PATTERN_CACHE.get(skill)
    if pattern is None:
        pattern = re.compile(
            r"(?<![a-zA-Z0-9])" + re.escape(skill) + r"(?![a-zA-Z0-9])",
            re.IGNORECASE,
        )
        _SKILL_PATTERN_CACHE[skill] = pattern
    return pattern.search(text) is not None


# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------
def extract_text_from_bytes(data: bytes, filename: str) -> str:
    """Extract raw text from a PDF, DOCX or TXT byte payload."""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext == "pdf":
        return _extract_pdf(data)
    if ext == "docx":
        return _extract_docx(data)
    if ext in ("txt", "md", ""):
        return data.decode("utf-8", errors="ignore")
    raise ValueError(f"Unsupported file type: {ext or 'unknown'}. Use PDF, DOCX, or TXT.")


def _extract_pdf(data: bytes) -> str:
    import fitz  # PyMuPDF

    doc = fitz.open(stream=data, filetype="pdf")
    try:
        pages = [page.get_text("text") for page in doc]
    finally:
        doc.close()
    return "\n".join(pages).strip()


def _extract_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    lines: List[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    lines.append(cell.text.strip())
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Heuristic structured parsing
# ---------------------------------------------------------------------------
def _header_category(line: str) -> Optional[str]:
    stripped = line.strip()
    normalized = stripped.strip(":-•*·_ ").strip().lower()
    if normalized in SECTION_ALIASES:
        return SECTION_ALIASES[normalized]
    if re.fullmatch(r"[A-Z][A-Z &/()+\-\.]{2,}", stripped):
        lowered = stripped.lower().rstrip(".:- ").strip()
        if lowered in SECTION_ALIASES:
            return SECTION_ALIASES[lowered]
    return None


def _split_sections(text: str):
    """Split resume text into (category, body) sections."""
    sections: List[tuple] = []
    current: List[str] = []
    current_cat = "header"

    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            current.append("")
            continue
        category = _header_category(stripped)
        if category is not None:
            if current:
                sections.append((current_cat, current))
            current = []
            current_cat = category
        else:
            current.append(line)

    if current:
        sections.append((current_cat, current))
    return sections


def _guess_name(text: str) -> str:
    first_lines = [l.strip() for l in text.splitlines() if l.strip()]
    for line in first_lines[:5]:
        if line.lower().startswith(("name:", "name :")):
            return line.split(":", 1)[1].strip().title()
        words = line.split()
        if (
            1 <= len(words) <= 4
            and all(w.isalpha() for w in words)
            and line.lower().strip() not in NAME_STOPWORDS
            and not re.search(r"\d", line)
        ):
            return line.title()
    return "Not Found"


def _extract_education(body: List[str]) -> List[Education]:
    results: List[Education] = []
    for line in body:
        line = line.strip().lstrip("-•*·").strip()
        if not line or not any(k in line.lower() for k in EDUCATION_KEYWORDS):
            continue
        degree_match = re.search(
            r"((?:bachelor|master|ph\.?d|mba|b\.tech|m\.tech|b\.e|m\.e|m\.sc|b\.sc|"
            r"m\.c\.a|b\.c\.a|associate|diploma)[^,;|]*?)(?:degree)?",
            line,
            re.IGNORECASE,
        )
        degree = degree_match.group(1).strip() if degree_match else line
        year_match = re.search(r"(19|20)\d{2}", line)
        results.append(
            Education(
                degree=degree[:120],
                institution=line[:160],
                year=year_match.group(0) if year_match else None,
            )
        )
        if len(results) >= 4:
            break
    return results


def _extract_skills(lines: List[str]) -> List[str]:
    found: List[str] = []
    text = " ".join(lines)
    for skill in SKILL_LEXICON:
        if _matches(text, skill) and skill not in found:
            found.append(skill)
    priority = ["python", "machine learning", "tensorflow", "pytorch", "aws",
                "docker", "kubernetes", "sql", "java", "react", "fastapi"]
    found.sort(key=lambda s: (priority.index(s) if s in priority else 99, s))
    return found[:30]


def _years_from_duration(duration: Optional[str]) -> Optional[float]:
    if not duration:
        return None
    years = re.findall(r"(?:19|20)\d{2}", duration)
    is_current = re.search(r"present|current|now", duration, re.IGNORECASE)
    if years and is_current:
        return max(0.0, float(datetime.now().year) - int(years[0]))
    if len(years) >= 2:
        return max(0.0, float(years[1]) - float(years[0]))
    return None


def _extract_years_from_line(line: str) -> Optional[float]:
    match = YEARS_RE.search(line.lower())
    if match:
        return float(match.group(1))
    return None


def _looks_like_company_line(line: str) -> bool:
    return bool(
        re.search(
            r"(at |@ |llc|ltd|inc|corp|pvt|technolog|systems|solutions|"
            r"company|lab|studios|academy|university|firm)",
            line,
            re.IGNORECASE,
        )
    )


def _looks_like_role_line(line: str) -> bool:
    if len(line) > 70 or line.endswith("."):
        return False
    return bool(
        re.search(
            r"(engineer|developer|analyst|scientist|intern|lead|manager|"
            r"architect|consultant|designer|researcher|specialist)",
            line,
            re.IGNORECASE,
        )
    ) and not _looks_like_company_line(line)


def _extract_experience(body: List[str]) -> List[Experience]:
    experiences: List[Experience] = []
    current: Optional[Experience] = None

    for raw in body:
        line = raw.strip().lstrip("-•*·").strip()
        if not line:
            continue

        if DATE_RANGE_RE.match(line):
            if current is None:
                current = Experience(role="Not Found", company="Not Found")
                experiences.append(current)
            current.duration = line
            current.years = _years_from_duration(line)
            continue

        if "|" in line:
            parts = [p.strip() for p in line.split("|")]
            role = parts[0] or "Not Found"
            company = parts[1] if len(parts) > 1 and parts[1] else "Not Found"
            current = Experience(
                role=role[:120],
                company=company[:120],
                years=_extract_years_from_line(line),
            )
            experiences.append(current)
            continue

        if _looks_like_role_line(line):
            current = Experience(
                role=line[:120],
                company="Not Found",
                years=_extract_years_from_line(line),
            )
            experiences.append(current)
            continue

        if current is None:
            continue
        if _looks_like_company_line(line):
            current.company = line[:120]
        else:
            current.description.append(line[:200])

    for exp in experiences:
        if exp.years is None:
            exp.years = (
                _extract_years_from_line(" ".join(exp.description))
                or _years_from_duration(exp.duration)
            )

    return experiences[:6]


def _is_project_heading(line: str) -> bool:
    if " - " in line or " – " in line:
        return True
    if len(line) > 80:
        return False
    if line.startswith(
        ("Built", "Developed", "Created", "Designed", "Led", "Implemented", "Engineered")
    ):
        return True
    if line.endswith((":", ".")):
        return True
    if re.fullmatch(r"[A-Za-z0-9 .,&'()\-]+", line) and len(line.split()) <= 6:
        return True
    return False


def _extract_projects(body: List[str]) -> List[Project]:
    projects: List[Project] = []
    current: Optional[Project] = None

    for raw in body:
        line = raw.strip().lstrip("-•*·").strip()
        if not line:
            continue
        if _is_project_heading(line):
            if " - " in line:
                name, _, desc = line.partition(" - ")
            elif " – " in line:
                name, _, desc = line.partition(" – ")
            else:
                name, desc = line, ""
            current = Project(name=name.strip()[:120], description=desc.strip()[:200])
            projects.append(current)
            continue
        if current is not None:
            if not current.description:
                current.description = line[:200]
            else:
                current.description += " " + line[:200]

    for proj in projects:
        proj.skills = [
            s for s in SKILL_LEXICON if _matches(f"{proj.name} {proj.description}", s)
        ][:6]

    return projects[:6]


def _extract_certifications(body: List[str]) -> List[str]:
    certs: List[str] = []
    for line in body:
        line = line.strip().lstrip("-•*·").strip()
        if any(k in line.lower() for k in CERTIFICATION_KEYWORDS):
            certs.append(line[:160])
    return list(dict.fromkeys(certs))[:8]


def _extract_achievements(body: List[str]) -> List[str]:
    return [l.strip().lstrip("-•*·").strip()[:200] for l in body if l.strip()][:8]


def _extract_summary(body: List[str]) -> Optional[str]:
    lines = [l.strip().lstrip("-•*·").strip() for l in body if l.strip()]
    if not lines:
        return None
    text = " ".join(lines)
    # Trim if the "summary" section was actually the header/contact block.
    if len(text) < 120:
        return None
    return text[:600]


def parse_resume_text(text: str) -> ParsedResume:
    """Run the heuristic structure pass over extracted text."""
    email_match = EMAIL_RE.search(text)
    phone_match = PHONE_RE.search(text)

    education: List[Education] = []
    skills: List[str] = []
    experience: List[Experience] = []
    projects: List[Project] = []
    certifications: List[str] = []
    achievements: List[str] = []
    summary: Optional[str] = None

    for category, body in _split_sections(text):
        if category == "education":
            education += _extract_education(body)
        elif category == "skills":
            skills += _extract_skills(body)
        elif category == "experience":
            experience += _extract_experience(body)
        elif category == "projects":
            projects += _extract_projects(body)
        elif category == "certifications":
            certifications += _extract_certifications(body)
        elif category == "achievements":
            achievements += _extract_achievements(body)
        elif category == "summary" and summary is None:
            summary = _extract_summary(body)

    if not skills:
        skills = _extract_skills(text.splitlines())

    return ParsedResume(
        name=_guess_name(text),
        email=email_match.group(0) if email_match else "Not Found",
        phone=phone_match.group(0) if phone_match else "Not Found",
        summary=summary,
        education=education,
        skills=skills,
        workExperience=experience,
        projects=projects,
        certifications=certifications,
        achievements=achievements,
    )


def parse_resume_bytes(data: bytes, filename: str) -> ParsedResume:
    """Extract + structure a resume from its raw bytes."""
    text = extract_text_from_bytes(data, filename)
    return parse_resume_text(text)
